from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOC_PATH = Path(__file__).with_name("User_Manual_HVAS_expanded.docx")


def clear_paragraph(paragraph):
    paragraph.text = ""


def set_paragraph_text(paragraph, text, *, bold=False, size=None, color=None):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return paragraph


def set_cell_text(cell, text, *, bold=False, size=9, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="B7C4D6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def format_table(table, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(width)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if row_idx == 0:
                shade_cell(cell, "2F5597")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_idx % 2 == 0:
                shade_cell(cell, "F3F6FA")


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    set_paragraph_text(paragraph, text, bold=True, size=14 if level == 1 else 11, color=(31, 78, 121))
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_text(paragraph, text, size=9)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph()
    try:
        paragraph.style = "List Bullet"
    except Exception:
        pass
    set_paragraph_text(paragraph, text, size=9)
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=(255, 255, 255))
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value)
    format_table(table, widths=widths)
    return table


def move_blocks_before(blocks, reference_paragraph):
    for block in blocks:
        reference_paragraph._p.addprevious(block)


def update_document_control(doc):
    if not doc.tables:
        return
    table = doc.tables[0]
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = row.cells[0].text.strip().lower()
        if label == "version":
            set_cell_text(row.cells[1], "1.1")
        elif label == "date":
            set_cell_text(row.cells[1], "2026-05-06")
        elif label == "status":
            set_cell_text(row.cells[1], "Final portfolio version - deployment guidance expanded")
        elif label == "source basis":
            set_cell_text(
                row.cells[1],
                "System source code, frontend screens, backend API routes, docker-compose.yml, "
                "requirements.txt, frontend/package.json, and .env.example.",
            )


def renumber_sections(doc):
    replacements = {
        "2. Main Workflows": "3. Main Workflows",
        "3. Screenshot Evidence Fields": "4. Screenshot Evidence Fields",
        "4. Troubleshooting for Users": "5. Troubleshooting for Users",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text], bold=True, size=14, color=(31, 78, 121))
            paragraph.paragraph_format.keep_with_next = True


def apply_section_page_breaks(doc):
    break_before = {
        "2.2 Install Project Dependencies",
        "2.3 Configure .env from .env.example",
        "2.4 Start Infrastructure with Docker Desktop",
        "3. Main Workflows",
        "5. Troubleshooting for Users",
    }
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in break_before:
            paragraph.paragraph_format.page_break_before = True


def append_troubleshooting_rows(doc):
    table = None
    for candidate in doc.tables:
        if not candidate.rows:
            continue
        headers = [cell.text.strip().lower() for cell in candidate.rows[0].cells]
        if headers and headers[0] == "problem":
            table = candidate
            break
    if table is None:
        return
    additions = [
        (
            "Docker command fails or MongoDB/MinIO stay disconnected",
            "Docker Desktop is not running, the Docker engine is still starting, or ports 27017/9000/9001 are occupied.",
            "Start Docker Desktop, wait until it reports that the engine is running, then execute docker compose up -d and docker compose ps from the project root.",
        ),
        (
            "python or pip is not recognized",
            "Python was installed without PATH registration, or the terminal was opened before installation completed.",
            "Reopen the terminal, run python --version and pip --version, or reinstall Python from python.org and select Add python.exe to PATH.",
        ),
        (
            "Backend starts but /api/health returns degraded",
            "The backend is reachable, but MongoDB or MinIO cannot be reached.",
            "Confirm docker compose ps shows mongo and minio as running, then check MONGO_URI, MINIO_ENDPOINT, MINIO_ACCESS_KEY, and MINIO_SECRET_KEY in .env.",
        ),
        (
            "Qwen report generation or vision verification is unavailable",
            "The relevant Qwen API key, endpoint, or model name is empty or incorrect.",
            "Copy .env.example to .env, set QWEN_REPORT_API_KEY for reports and QWEN_VL_API_KEY if Qwen-VL verification is used, then restart the backend.",
        ),
        (
            "Webhook receiver rejects HVAS events",
            "The HMAC secret configured in HVAS does not match the secret configured in the receiving system.",
            "Set WEBHOOK_SECRET to a shared production value in .env and configure the same value in the receiver before registering webhooks.",
        ),
    ]
    for issue, cause, action in additions:
        cells = table.add_row().cells
        if len(cells) == 2:
            values = (issue, f"Cause: {cause}\nAction: {action}")
        else:
            values = (issue, cause, action)
        for idx, value in enumerate(values):
            if idx < len(cells):
                set_cell_text(cells[idx], value)
    format_table(table)


def build_deployment_section(doc):
    blocks = []

    def capture(element):
        blocks.append(element._p if hasattr(element, "_p") else element._tbl)
        return element

    capture(add_heading(doc, "2. Deployment and Initial Setup", level=1))
    capture(
        add_body(
            doc,
            "This section describes a clean Windows deployment path for HVAS. The steps assume the operator "
            "starts from the project root and uses Docker Desktop for MongoDB and MinIO, Python for the Flask "
            "backend, and Node.js/npm for the Vue dashboard.",
        )
    )

    capture(add_heading(doc, "2.1 Download and Install Prerequisites", level=2))
    capture(
        add_table(
            doc,
            ["Component", "Where to Get It", "Installation Notes", "Verification"],
            [
                (
                    "Docker Desktop",
                    "https://www.docker.com/products/docker-desktop/",
                    "Install Docker Desktop for Windows. Enable WSL 2/virtualization if the installer prompts for it. Start Docker Desktop before running infrastructure commands.",
                    "docker --version\ndocker compose version",
                ),
                (
                    "Python",
                    "https://www.python.org/downloads/windows/",
                    "Install a project-compatible Python version. Select Add python.exe to PATH during installation, then reopen the terminal.",
                    "python --version\npip --version",
                ),
                (
                    "Node.js and npm",
                    "https://nodejs.org/",
                    "Install the LTS release if npm is not already available. The Vue dashboard and upload service use npm scripts from frontend/package.json.",
                    "node --version\nnpm --version",
                ),
                (
                    "Git or project archive",
                    "Existing repository or delivered zip package",
                    "Use Git clone or extract the delivered archive. Open PowerShell in the project root before running the following commands.",
                    "dir docker-compose.yml\ndir .env.example",
                ),
            ],
            widths=[1.3, 1.9, 3.1, 1.6],
        )
    )

    capture(add_heading(doc, "2.2 Install Project Dependencies", level=2))
    capture(add_body(doc, "Run the backend dependency installation from the project root:"))
    capture(
        add_table(
            doc,
            ["Step", "PowerShell Command", "Purpose"],
            [
                ("Create virtual environment", "python -m venv venv", "Keeps project Python packages isolated."),
                ("Activate virtual environment", ".\\venv\\Scripts\\activate", "Ensures pip installs into the project environment."),
                ("Upgrade pip", "python -m pip install --upgrade pip", "Reduces dependency installation failures."),
                ("Install backend packages", "pip install -r requirements.txt", "Installs Flask, OpenCV, Ultralytics YOLO, MongoDB, MinIO, requests, dotenv, and watchdog dependencies."),
                ("Optional CUDA torch", "pip install torch torchvision --index-url https://download.pytorch.org/whl/cu121", "Use only when the deployment machine has a compatible NVIDIA GPU and CUDA runtime."),
                ("Install frontend packages", "cd frontend\nnpm install\ncd ..", "Installs Vue, Vite, axios, three/globe.gl, Express, MongoDB driver, and multer."),
            ],
            widths=[1.5, 2.6, 3.4],
        )
    )

    capture(add_heading(doc, "2.3 Configure .env from .env.example", level=2))
    capture(
        add_body(
            doc,
            "Create a local .env file from the template and fill in deployment-specific API keys and secrets. "
            "Do not commit .env to source control.",
        )
    )
    capture(
        add_table(
            doc,
            ["Step", "PowerShell Command or Setting", "Notes"],
            [
                ("Copy template", "Copy-Item .env.example .env", "Use the template as the starting point for all local configuration."),
                ("Set webhook secret", "WEBHOOK_SECRET=<strong-shared-secret>", "Required when ENABLE_WEBHOOKS=true. The receiving MUBS/webhook system must use the same HMAC secret."),
                ("Set Qwen report key", "QWEN_REPORT_API_KEY=<your DashScope key>", "Required for LLM-based common-space reports. QWEN_REPORT_API_URL defaults to the DashScope OpenAI-compatible endpoint in .env.example."),
                ("Set Qwen-VL key if used", "QWEN_VL_API_KEY=<your DashScope key>", "Needed when smoke/fire secondary verification or visual common-space analysis depends on Qwen-VL environment configuration."),
                ("Choose inference device", "YOLO_DEVICE=cuda or YOLO_DEVICE=cpu", "Use cuda only after confirming torch.cuda.is_available() on the deployment machine."),
                ("Review RTSP sampling", "RTSP_SAMPLE_INTERVAL=0.5", "Default legacy sampling interval. The backend also supports PARKING_RTSP_SAMPLE_INTERVAL, SMOKE_RTSP_SAMPLE_INTERVAL, and COMMON_SPACE_RTSP_SAMPLE_INTERVAL for task-level overrides."),
                ("Review queue limits", "STREAM_EXECUTOR_MAX_QUEUE=24\nSMOKE_DETECTION_MAX_QUEUE=8\nSMOKE_VERIFICATION_MAX_QUEUE=12\nCOMMON_SPACE_MAX_QUEUE=4", "Keep backpressure enabled for multi-stream deployments to avoid unbounded queue growth."),
            ],
            widths=[1.45, 2.75, 3.25],
        )
    )
    capture(add_bullet(doc, "Docker Compose provides default local MongoDB and MinIO settings: MONGO_URI=mongodb://localhost:27017, MONGO_DB_NAME=video_analysis_db, MINIO_ENDPOINT=localhost:9000, MINIO_ACCESS_KEY=minioadmin, MINIO_SECRET_KEY=minioadmin, and MINIO_BUCKET=video-events."))
    capture(add_bullet(doc, "For production or shared demos, replace default MinIO credentials and WEBHOOK_SECRET with non-default values, then restart the affected services."))

    capture(add_heading(doc, "2.4 Start Infrastructure with Docker Desktop", level=2))
    capture(
        add_table(
            doc,
            ["Step", "PowerShell Command", "Expected Result"],
            [
                ("Start Docker Desktop", "Open Docker Desktop from Windows Start Menu", "Docker engine is running."),
                ("Start MongoDB and MinIO", "docker compose up -d", "Containers named mongo and minio are created or restarted."),
                ("Check container status", "docker compose ps", "MongoDB exposes 27017. MinIO exposes API 9000 and console 9001."),
                ("Open MinIO console if needed", "http://localhost:9001", "Login with configured MinIO credentials. Local defaults are minioadmin / minioadmin."),
            ],
            widths=[1.5, 2.6, 3.4],
        )
    )

    capture(add_heading(doc, "2.5 Start HVAS Services", level=2))
    capture(
        add_table(
            doc,
            ["Service", "Command", "URL or Output"],
            [
                ("Backend Flask API", "python -m backend.main", "http://localhost:5000"),
                ("Frontend dashboard", "cd frontend\nnpm run dev", "http://localhost:5173"),
                ("Upload/proxy service", "cd frontend\nnpm run start-server", "http://localhost:8080"),
            ],
            widths=[1.6, 3.0, 2.6],
        )
    )
    capture(add_body(doc, "The delivered batch launcher can also be used after dependencies are installed: run the project-root one-click launcher, then verify each service window remains open without errors."))

    capture(add_heading(doc, "2.6 Post-Deployment Verification", level=2))
    capture(
        add_table(
            doc,
            ["Check", "Command or UI Action", "Pass Criteria"],
            [
                ("Backend health", "Invoke-RestMethod http://localhost:5000/api/health", "status is ok or degraded only when optional services are intentionally unavailable; mongodb and minio should be connected for a full deployment."),
                ("Dashboard access", "Open http://localhost:5173", "Dashboard, Stream Manager, Events, Globe, and Reports pages load."),
                ("RTSP stream setup", "Add a stream in Stream Manager", "The stream appears as running and /api/streams returns the configured camera_id and task list."),
                ("Event persistence", "Trigger a test event, then open Events List", "New event card appears and the event image is reachable from MinIO."),
                ("Report generation", "Open Reports and run common-space generation", "Rule-based reports work; LLM-based reports require QWEN_REPORT_API_KEY."),
            ],
            widths=[1.4, 2.6, 3.6],
        )
    )

    capture(add_heading(doc, "2.7 Shutdown and Restart", level=2))
    capture(add_bullet(doc, "Stop backend and frontend terminals with Ctrl+C."))
    capture(add_bullet(doc, "Stop infrastructure with docker compose down when local data should be preserved in Docker volumes."))
    capture(add_bullet(doc, "After changing .env, restart the backend so python-dotenv reloads the updated settings."))

    return blocks


def main():
    doc = Document(DOC_PATH)
    update_document_control(doc)

    reference = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "2. Main Workflows":
            reference = paragraph
            break
    if reference is None:
        raise RuntimeError("Could not find insertion point: 2. Main Workflows")

    blocks = build_deployment_section(doc)
    move_blocks_before(blocks, reference)
    renumber_sections(doc)
    apply_section_page_breaks(doc)
    append_troubleshooting_rows(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    doc.save(DOC_PATH)


if __name__ == "__main__":
    main()
